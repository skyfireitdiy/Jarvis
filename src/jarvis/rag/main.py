import os
import hashlib
import numpy as np
import faiss
from typing import List, Tuple, Optional, Dict
from sentence_transformers import SentenceTransformer
import pickle
from jarvis.utils import OutputType, PrettyOutput, find_git_root, load_embedding_model
from jarvis.utils import load_env_from_file
import tiktoken
from dataclasses import dataclass
from tqdm import tqdm
import fitz  # PyMuPDF for PDF files
from docx import Document as DocxDocument  # python-docx for DOCX files
from pathlib import Path
from jarvis.models.registry import PlatformRegistry

@dataclass
class Document:
    """文档类，用于存储文档内容和元数据"""
    content: str  # 文档内容
    metadata: Dict  # 元数据(文件路径、位置等)

class FileProcessor:
    """文件处理器基类"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        """判断是否可以处理该文件"""
        raise NotImplementedError
        
    @staticmethod
    def extract_text(file_path: str) -> str:
        """提取文件文本内容"""
        raise NotImplementedError

class TextFileProcessor(FileProcessor):
    """文本文件处理器"""
    ENCODINGS = ['utf-8', 'gbk', 'gb2312', 'latin1']
    SAMPLE_SIZE = 8192  # 读取前8KB来检测编码
    
    @staticmethod
    def can_handle(file_path: str) -> bool:
        """判断文件是否为文本文件，通过尝试解码来判断"""
        try:
            # 读取文件开头的一小部分来检测
            with open(file_path, 'rb') as f:
                sample = f.read(TextFileProcessor.SAMPLE_SIZE)
                
            # 检查是否包含空字节（通常表示二进制文件）
            if b'\x00' in sample:
                return False
                
            # 检查是否包含过多的非打印字符（通常表示二进制文件）
            non_printable = sum(1 for byte in sample if byte < 32 and byte not in (9, 10, 13))  # tab, newline, carriage return
            if non_printable / len(sample) > 0.3:  # 如果非打印字符超过30%，认为是二进制文件
                return False
                
            # 尝试用不同编码解码
            for encoding in TextFileProcessor.ENCODINGS:
                try:
                    sample.decode(encoding)
                    return True
                except UnicodeDecodeError:
                    continue
                    
            return False
            
        except Exception:
            return False
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """提取文本内容，使用检测到的正确编码"""
        detected_encoding = None
        try:
            # 首先尝试检测编码
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                
            # 尝试不同的编码
            for encoding in TextFileProcessor.ENCODINGS:
                try:
                    raw_data.decode(encoding)
                    detected_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
                    
            if not detected_encoding:
                raise UnicodeDecodeError(f"无法用支持的编码解码文件: {file_path}")
                
            # 使用检测到的编码读取文件
            with open(file_path, 'r', encoding=detected_encoding, errors='replace') as f:
                content = f.read()
                
            # 规范化Unicode字符
            import unicodedata
            content = unicodedata.normalize('NFKC', content)
            
            return content
            
        except Exception as e:
            raise Exception(f"读取文件失败: {str(e)}")

class PDFProcessor(FileProcessor):
    """PDF文件处理器"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.pdf'
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        text_parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        return "\n".join(text_parts)

class DocxProcessor(FileProcessor):
    """DOCX文件处理器"""
    @staticmethod
    def can_handle(file_path: str) -> bool:
        return Path(file_path).suffix.lower() == '.docx'
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        doc = DocxDocument(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

class RAGTool:
    def __init__(self, root_dir: str):
        """初始化RAG工具
        
        Args:
            root_dir: 项目根目录
        """
        load_env_from_file()
        self.root_dir = root_dir
        os.chdir(self.root_dir)
        
        # 初始化配置
        self.min_paragraph_length = int(os.environ.get("JARVIS_MIN_PARAGRAPH_LENGTH", "50"))  # 最小段落长度
        self.max_paragraph_length = int(os.environ.get("JARVIS_MAX_PARAGRAPH_LENGTH", "1000"))  # 最大段落长度
        
        # 初始化数据目录
        self.data_dir = os.path.join(self.root_dir, ".jarvis-rag")
        if not os.path.exists(self.data_dir):
            os.makedirs(self.data_dir)
            
        # 初始化嵌入模型
        try:
            self.embedding_model = load_embedding_model()
            self.vector_dim = self.embedding_model.get_sentence_embedding_dimension()
            PrettyOutput.print("模型加载完成", output_type=OutputType.SUCCESS)
        except Exception as e:
            PrettyOutput.print(f"加载模型失败: {str(e)}", output_type=OutputType.ERROR)
            raise

        # 初始化缓存和索引
        self.cache_path = os.path.join(self.data_dir, "cache.pkl")
        self.documents: List[Document] = []
        self.index = None
        
        # 加载缓存
        self._load_cache()

        # 注册文件处理器
        self.file_processors = [
            TextFileProcessor(),
            PDFProcessor(),
            DocxProcessor()
        ]

    def _load_cache(self):
        """加载缓存数据"""
        if os.path.exists(self.cache_path):
            try:
                with open(self.cache_path, 'rb') as f:
                    cache_data = pickle.load(f)
                    self.documents = cache_data["documents"]
                    vectors = cache_data["vectors"]
                    
                # 重建索引
                self._build_index(vectors)
                PrettyOutput.print(f"加载了 {len(self.documents)} 个文档片段", 
                                output_type=OutputType.INFO)
            except Exception as e:
                PrettyOutput.print(f"加载缓存失败: {str(e)}", 
                                output_type=OutputType.WARNING)
                self.documents = []
                self.index = None

    def _save_cache(self, vectors: np.ndarray):
        """保存缓存数据"""
        try:
            cache_data = {
                "documents": self.documents,
                "vectors": vectors
            }
            with open(self.cache_path, 'wb') as f:
                pickle.dump(cache_data, f)
            PrettyOutput.print(f"保存了 {len(self.documents)} 个文档片段", 
                            output_type=OutputType.INFO)
        except Exception as e:
            PrettyOutput.print(f"保存缓存失败: {str(e)}", 
                            output_type=OutputType.ERROR)

    def _build_index(self, vectors: np.ndarray):
        """构建FAISS索引"""
        # 创建HNSW索引
        hnsw_index = faiss.IndexHNSWFlat(self.vector_dim, 16)
        hnsw_index.hnsw.efConstruction = 40
        hnsw_index.hnsw.efSearch = 16
        
        # 用IndexIDMap包装HNSW索引
        self.index = faiss.IndexIDMap(hnsw_index)
        
        # 添加向量到索引
        if vectors.shape[0] > 0:
            self.index.add_with_ids(vectors, np.arange(vectors.shape[0]))
        else:
            self.index = None

    def _split_text(self, text: str) -> List[str]:
        """将文本分割成段落
        
        Args:
            text: 要分割的文本
            
        Returns:
            分割后的段落列表
        """
        # 首先按空行分割
        paragraphs = []
        current_paragraph = []
        
        for line in text.split('\n'):
            line = line.strip()
            if not line:  # 空行表示段落结束
                if current_paragraph:
                    paragraph_text = ' '.join(current_paragraph)
                    if len(paragraph_text) >= self.min_paragraph_length:
                        paragraphs.append(paragraph_text)
                    current_paragraph = []
            else:
                current_paragraph.append(line)
        
        # 处理最后一个段落
        if current_paragraph:
            paragraph_text = ' '.join(current_paragraph)
            if len(paragraph_text) >= self.min_paragraph_length:
                paragraphs.append(paragraph_text)
        
        # 处理过长的段落
        final_paragraphs = []
        for paragraph in paragraphs:
            if len(paragraph) <= self.max_paragraph_length:
                final_paragraphs.append(paragraph)
            else:
                # 按句子分割过长的段落
                sentences = []
                current_sentence = []
                
                # 中文句子结束标记
                sentence_ends = {'。', '！', '？', '…', '.', '!', '?'}
                
                for char in paragraph:
                    current_sentence.append(char)
                    if char in sentence_ends:
                        sentence = ''.join(current_sentence)
                        if sentence.strip():
                            sentences.append(sentence)
                        current_sentence = []
                
                # 处理最后一个句子
                if current_sentence:
                    sentence = ''.join(current_sentence)
                    if sentence.strip():
                        sentences.append(sentence)
                
                # 组合句子成适当长度的段落
                current_chunk = []
                current_length = 0
                
                for sentence in sentences:
                    sentence_length = len(sentence)
                    if current_length + sentence_length > self.max_paragraph_length:
                        if current_chunk:
                            final_paragraphs.append(''.join(current_chunk))
                        current_chunk = [sentence]
                        current_length = sentence_length
                    else:
                        current_chunk.append(sentence)
                        current_length += sentence_length
                
                # 处理最后一个chunk
                if current_chunk:
                    final_paragraphs.append(''.join(current_chunk))
        
        # 过滤掉太短的段落
        final_paragraphs = [p for p in final_paragraphs if len(p) >= self.min_paragraph_length]
        
        return final_paragraphs

    def _get_embedding(self, text: str) -> np.ndarray:
        """获取文本的向量表示"""
        embedding = self.embedding_model.encode(text, 
                                            normalize_embeddings=True,
                                            show_progress_bar=False)
        return np.array(embedding, dtype=np.float32)

    def _process_file(self, file_path: str) -> List[Document]:
        """处理单个文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档对象列表
        """
        try:
            # 查找合适的处理器
            processor = None
            for p in self.file_processors:
                if p.can_handle(file_path):
                    processor = p
                    break
                    
            if not processor:
                PrettyOutput.print(f"跳过不支持的文件: {file_path}", 
                                output_type=OutputType.WARNING)
                return []
            
            # 提取文本内容
            content = processor.extract_text(file_path)
            if not content.strip():
                PrettyOutput.print(f"文件内容为空: {file_path}", 
                                output_type=OutputType.WARNING)
                return []
            
                
            # 分割文本
            chunks = self._split_text(content)
            
            # 创建文档对象
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    content=chunk,
                    metadata={
                        "file_path": file_path,
                        "file_type": Path(file_path).suffix.lower(),
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    }
                )
                documents.append(doc)
                
            return documents
            
        except Exception as e:
            PrettyOutput.print(f"处理文件失败 {file_path}: {str(e)}", 
                            output_type=OutputType.ERROR)
            return []

    def build_index(self, dir: str):
        """构建文档索引"""
        # 获取所有文件
        all_files = []
        for root, _, files in os.walk(dir):
            if any(ignored in root for ignored in ['.jarvis-rag', '.git', '__pycache__', 'node_modules']):
                continue
            for file in files:
                file_path = os.path.join(root, file)
                # 跳过大文件
                if os.path.getsize(file_path) > 10 * 1024 * 1024:  # 10MB
                    PrettyOutput.print(f"跳过大文件: {file_path}", 
                                    output_type=OutputType.WARNING)
                    continue
                all_files.append(file_path)

        # 处理所有文件
        self.documents = []
        for file_path in tqdm(all_files, desc="处理文件"):
            docs = self._process_file(file_path)
            self.documents.extend(docs)

        # 获取所有文档的向量表示
        vectors = []
        for doc in tqdm(self.documents, desc="生成向量"):
            vector = self._get_embedding(doc.content)
            vectors.append(vector)

        if vectors:
            vectors = np.vstack(vectors)
            # 构建索引
            self._build_index(vectors)
            # 保存缓存
            self._save_cache(vectors)
            
        PrettyOutput.print(f"成功索引了 {len(self.documents)} 个文档片段", 
                        output_type=OutputType.SUCCESS)

    def search(self, query: str, top_k: int = 5) -> List[Tuple[Document, float]]:
        """搜索相关文档
        
        Args:
            query: 查询文本
            top_k: 返回结果数量
            
        Returns:
            文档和相似度得分的列表
        """
        if not self.index:
            PrettyOutput.print("索引未构建，正在构建...", output_type=OutputType.INFO)
            self.build_index(self.root_dir)
            
        # 获取查询的向量表示
        query_vector = self._get_embedding(query)
        query_vector = query_vector.reshape(1, -1)
        
        # 搜索最相似的向量
        distances, indices = self.index.search(query_vector, top_k)
        
        # 返回结果
        results = []
        for idx, distance in zip(indices[0], distances[0]):
            if idx == -1:  # FAISS返回-1表示无效结果
                continue
            similarity = 1.0 / (1.0 + float(distance))
            results.append((self.documents[idx], similarity))
                
        return results

    def is_index_built(self):
        """检查索引是否已构建"""
        return self.index is not None

    def query(self, query: str) -> List[Document]:
        """查询相关文档
        
        Args:
            query: 查询文本
            
        Returns:
            相关文档列表
        """
        if not self.is_index_built():
            raise ValueError("索引未构建，请先调用build_index()")
            
        results = self.search(query)
        return [doc for doc, _ in results]

    def ask(self, question: str) -> Optional[str]:
        """询问关于文档的问题
        
        Args:
            question: 用户问题
            
        Returns:
            模型回答，如果失败则返回 None
        """
        try:
            # 搜索相关文档片段
            results = self.query(question)
            if not results:
                return None

            # 构建上下文
            context = []
            for doc in results:
                context.append(f"""
来源文件: {doc.metadata['file_path']}
片段位置: {doc.metadata['chunk_index'] + 1}/{doc.metadata['total_chunks']}
内容:
{doc.content}
---
""")

            # 构建提示词
            prompt = f"""请基于以下文档片段回答用户的问题。如果文档片段中的信息不足以完整回答问题，请明确指出。
            
用户问题: {question}

相关文档片段:
{''.join(context)}

请提供准确、简洁的回答，并在适当时引用具体的文档来源。
"""
            # 获取模型实例并生成回答
            model = PlatformRegistry.get_global_platform_registry().get_normal_platform()
            response = model.chat(prompt)
            
            return response
            
        except Exception as e:
            PrettyOutput.print(f"问答失败: {str(e)}", output_type=OutputType.ERROR)
            return None

def main():
    """主函数"""
    import argparse
    import sys
    
    # 设置标准输出编码为UTF-8
    if sys.stdout.encoding != 'utf-8':
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    
    parser = argparse.ArgumentParser(description='文档检索和分析工具')
    parser.add_argument('--dir', type=str, help='要处理的文档目录')
    parser.add_argument('--build', action='store_true', help='构建文档索引')
    parser.add_argument('--search', type=str, help='搜索文档内容')
    parser.add_argument('--ask', type=str, help='询问关于文档的问题')
    args = parser.parse_args()

    try:
        current_dir = find_git_root()
        rag = RAGTool(current_dir)

        if args.dir and args.build:
            PrettyOutput.print(f"正在处理目录: {args.dir}", output_type=OutputType.INFO)
            rag.build_index(args.dir)
            return 0

        if args.search or args.ask:
            if not rag.is_index_built():
                PrettyOutput.print("索引尚未构建，请先使用 --dir 和 --build 参数构建索引", output_type=OutputType.WARNING)
                return 1

            if args.search:
                results = rag.query(args.search)
                if not results:
                    PrettyOutput.print("未找到相关内容", output_type=OutputType.WARNING)
                    return 1
                    
                for doc in results:
                    PrettyOutput.print(f"\n文件: {doc.metadata['file_path']}", output_type=OutputType.INFO)
                    PrettyOutput.print(f"片段 {doc.metadata['chunk_index'] + 1}/{doc.metadata['total_chunks']}", 
                                    output_type=OutputType.INFO)
                    PrettyOutput.print("\n内容:", output_type=OutputType.INFO)
                    content = doc.content.encode('utf-8', errors='replace').decode('utf-8')
                    PrettyOutput.print(content, output_type=OutputType.INFO)
                return 0

            if args.ask:
                # 调用 ask 方法
                response = rag.ask(args.ask)
                if not response:
                    PrettyOutput.print("未能获取答案", output_type=OutputType.WARNING)
                    return 1
                    
                # 显示回答
                PrettyOutput.print("\n回答:", output_type=OutputType.INFO)
                PrettyOutput.print(response, output_type=OutputType.INFO)
                return 0

        PrettyOutput.print("请指定操作参数。使用 -h 查看帮助。", output_type=OutputType.WARNING)
        return 1

    except Exception as e:
        PrettyOutput.print(f"执行失败: {str(e)}", output_type=OutputType.ERROR)
        return 1

if __name__ == "__main__":
    main()
